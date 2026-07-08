"""Generates Borama_Hardware_Supabase_Dashboard_Steps.docx — a step-by-step
walkthrough for the 3 manual Supabase Dashboard configuration steps needed to
finish Phase 1 (Phone OTP, Twilio SMS provider, Auth Hook registration).

Twilio credentials are read from .env.local at generation time (never
hardcoded here) since this script is committed to version control while
.env.local and the generated .docx are gitignored."""

import os
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def load_env_local():
    env_path = Path(__file__).resolve().parent.parent / ".env.local"
    values = {}
    for line in env_path.read_text(encoding="utf-8").splitlines():
        m = re.match(r'^\s*([A-Z0-9_]+)\s*=\s*(.*)\s*$', line)
        if m:
            values[m.group(1)] = m.group(2).strip().strip('"')
    return values

env = load_env_local()
TWILIO_ACCOUNT_SID = env["TWILIO_ACCOUNT_SID"]
TWILIO_AUTH_TOKEN = env["TWILIO_AUTH_TOKEN"]
TWILIO_PHONE_NUMBER = env["TWILIO_PHONE_NUMBER"]

ORANGE = RGBColor(0xE0, 0x6B, 0x1F)
DARK = RGBColor(0x2B, 0x2B, 0x2B)
GREY = RGBColor(0x66, 0x66, 0x66)

doc = Document()

# base font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = DARK

def set_cell_shading(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def add_title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = ORANGE

def add_subtitle(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.color.rgb = GREY

def add_h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = ORANGE

def add_h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = DARK

def add_step(number, text, bold_bits=None):
    p = doc.add_paragraph(style='List Number')
    r = p.add_run(text)
    r.font.size = Pt(11)

def add_body(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(11)
    r.bold = bold
    r.italic = italic
    return p

def add_field_table(rows):
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Light Grid Accent 2'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for label, value in rows:
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = value
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(10.5)
        row.cells[0].paragraphs[0].runs[0].font.bold = True
        set_cell_shading(row.cells[0], "FBEBDD")
    doc.add_paragraph()

def add_note(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    r = p.add_run("Note: " + text)
    r.font.size = Pt(10)
    r.italic = True
    r.font.color.rgb = GREY

def add_link_line(label, url):
    p = doc.add_paragraph()
    r1 = p.add_run(label + ": ")
    r1.font.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(url)
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(0x10, 0x6C, 0xD6)
    r2.underline = True

add_title("Borama Hardware")
add_subtitle("Supabase Dashboard Configuration — Step-by-Step Walkthrough")
add_subtitle("Project: AI Hardware Retail Store  (ref: ebxitnzrozdzimaaqvcg)")
doc.add_paragraph()

add_body(
    "This document walks through the 3 remaining manual configuration steps needed "
    "in the Supabase Dashboard to finish Phase 1 (Authentication). These steps cannot "
    "be done via API/MCP — they require clicking through the Dashboard UI directly. "
    "Once all 3 are complete, tell Claude \"done\" and it will run a live end-to-end "
    "OTP test (send code, verify, confirm your profile is auto-created)."
)

# ---------------------------------------------------------------
add_h1("Step 1 — Enable the Phone Auth Provider")

add_link_line("Go to", "https://supabase.com/dashboard/project/ebxitnzrozdzimaaqvcg/auth/providers")

steps1 = [
    "On the page that loads, you'll see a list of authentication providers (Email, Phone, Google, GitHub, etc.).",
    "Find the row labeled \"Phone\" and click on it to expand its settings panel.",
    "Toggle the switch labeled \"Enable Phone provider\" (or \"Enable Sign in with Phone\") to the ON position.",
    "Scroll down inside that same expanded panel until you find a field called \"OTP Expiry\" (sometimes labeled \"Phone OTP expiry\" or shown in seconds).",
    "Set that value to 600 (this means the 6-digit code texted to the customer is valid for 10 minutes).",
    "Look for a checkbox or toggle called \"Confirm phone number on signup\" (or \"Enable phone confirmations\") and make sure it is turned ON. This is what triggers the profile auto-creation trigger once verified.",
    "Click the \"Save\" button at the bottom of the panel before moving on.",
]
for i, s in enumerate(steps1, 1):
    add_step(i, s)

add_note("If you don't see the OTP expiry field directly, it may be under an \"Advanced settings\" expandable section within the same Phone panel.")

# ---------------------------------------------------------------
add_h1("Step 2 — Configure Twilio as the SMS Provider")

add_body(
    "This is usually part of the same Phone provider panel from Step 1 — once Phone is enabled, "
    "it should reveal an additional section for choosing how OTP codes actually get sent."
)

steps2 = [
    "Still within the expanded \"Phone\" panel (or in a nearby \"SMS Provider\" section if it's separated on your screen), find a dropdown labeled \"SMS Provider\".",
    "Select \"Twilio\" from that dropdown.",
    "New fields will appear asking for your Twilio credentials. Fill them in exactly as shown in the table below.",
    "Click \"Save\" once all three fields are filled in.",
]
for i, s in enumerate(steps2, 1):
    add_step(i, s)

add_field_table([
    ("Twilio Account SID", TWILIO_ACCOUNT_SID),
    ("Twilio Auth Token", TWILIO_AUTH_TOKEN),
    ("Message Service SID / Phone Number", TWILIO_PHONE_NUMBER),
])

add_note(
    "Depending on your Supabase Dashboard version, the last field may be labeled \"Twilio Message Service SID\" "
    "(if you use a Messaging Service) OR \"Twilio Phone Number\" (if sending directly from the number). "
    "Since this is a single phone number without a Messaging Service configured, enter +13203993656 in whichever "
    "field is shown."
)

# ---------------------------------------------------------------
add_h1("Step 3 — Register the Custom Access Token Auth Hook")

add_link_line("Go to", "https://supabase.com/dashboard/project/ebxitnzrozdzimaaqvcg/auth/hooks")

steps3 = [
    "On the Hooks page, find the card/section titled \"Customize Access Token (Auth Hook)\".",
    "Click \"Add hook\" or \"Enable\" on that card.",
    "You'll be asked to choose the hook type — select \"Postgres Function\" (as opposed to \"HTTP\").",
    "A dropdown will appear listing functions available in your public schema — select custom_access_token_hook.",
    "Click \"Save\" / \"Enable\" to confirm.",
    "Confirm the hook now shows as \"Enabled\" with custom_access_token_hook listed next to it.",
]
for i, s in enumerate(steps3, 1):
    add_step(i, s)

add_note(
    "This is the function that injects the customer's role (customer/staff/admin) and location_id into their "
    "JWT on every login, which is what all the Row-Level Security policies rely on. Without this hook enabled, "
    "every user would be treated as a plain customer regardless of their actual role."
)

# ---------------------------------------------------------------
add_h1("After You've Completed All 3 Steps")

add_body(
    "Come back and tell Claude \"done\". It will then:"
)
for i, s in enumerate([
    "Trigger a real OTP send to a test phone number.",
    "Verify the OTP code.",
    "Confirm a row was auto-created in the profiles table with the correct role and location claims baked into the JWT.",
], 1):
    add_step(i, s)

add_body("That closes out Phase 1 completely, and Phase 2 (Core Infrastructure / API layer) can begin.", bold=True)

doc.save(r"C:\Users\Samatar.Samatarschool\Desktop\AI Hardware Retail Store\Borama_Hardware_Supabase_Dashboard_Steps.docx")
print("done")
